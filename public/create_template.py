from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Crear un nuevo documento
doc = Document()

# Configurar márgenes
sections = doc.sections
for section in sections:
    section.top_margin = Pt(50)
    section.bottom_margin = Pt(50)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

# ASUNTO - Centrado, Negrita
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ASUNTO: ASIGNACIÓN DE CARGA ACADÉMICA')
run.bold = True
run.font.size = Pt(11)

# Fecha - Centrado
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Teziutlán, Pue., {fecha}')
run.font.size = Pt(10.5)

# Espacio
doc.add_paragraph()

# Nombre del maestro - Centrado, Negrita
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('{teacherName}')
run.bold = True
run.font.size = Pt(12)

# Tipo de personal - Centrado, Negrita
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('{personalType} DE LA UPN-212 TEZIUTLÁN')
run.bold = True
run.font.size = Pt(10)

# PRESENTE - Centrado, Negrita
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PRESENTE')
run.bold = True
run.font.size = Pt(10)

# Espacio
doc.add_paragraph()

# Párrafo introductorio - Justificado
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
text = 'En virtud de reunir características académicas y profesionales específicas; se le ha asignado el siguiente curso en la '
run = p.add_run(text)
run.font.size = Pt(10)
run_bold = p.add_run('{licenciatura}')
run_bold.bold = True
run_bold.font.size = Pt(10)
run2 = p.add_run(', que corresponde al periodo {periodo}..')
run2.font.size = Pt(10)

# Tabla de horarios
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'

# Encabezado de tabla
hdr_cells = table.rows[0].cells
headers = ['CLAVE', 'ASIGNATURA', 'HORARIO Y DÍA', 'SEM']
for i, header in enumerate(headers):
    hdr_cells[i].text = header
    for paragraph in hdr_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(9)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Sombreado gris claro
    shading_elm = hdr_cells[i]._element.get_or_add_tcPr()
    shading = qn('w:shd')
    shade_element = shading_elm.find(shading)
    if shade_element is None:
        shade_element = document.createElement(shading)
        shading_elm.append(shade_element)
    shade_element.set(qn('w:fill'), 'DCDCDC')

# Filas de datos (usando bucle de docxtemplater)
row_cells = table.add_row().cells
row_cells[0].text = '{#horarios}{clave}{/horarios}'
row_cells[1].text = '{#horarios}{asignatura}{/horarios}'
row_cells[2].text = '{#horarios}{horario}{/horarios}'
row_cells[3].text = '{#horarios}{semestre}{/horarios}'

for cell in row_cells:
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(9)

# Espacio
doc.add_paragraph()

# Párrafo de cierre - Justificado
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
text = 'Deseándole el mayor de los éxitos en el desempeño de esta tarea, aprovecho la oportunidad para invitarle a que, en el ejercicio de sus funciones, ponga lo mejor de su esfuerzo y dedicación al servicio de la Universidad Pedagógica Nacional Unidad 212 Teziutlán, siguiendo las indicaciones institucionales, estableciendo comunicación permanente con su coordinador(a) y apoyando en las diversas actividades que fortalecen la formación de nuestros alumnos, así como la vida institucional de nuestra universidad.'
run = p.add_run(text)
run.font.size = Pt(10)

# Espacio
doc.add_paragraph()
doc.add_paragraph()

# ATENTAMENTE - Centrado, Negrita
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ATENTAMENTE')
run.bold = True
run.font.size = Pt(10)

# Slogan - Centrado, Negrita
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('"EDUCAR PARA TRANSFORMAR"')
run.bold = True
run.font.size = Pt(10)

# Espacio para firma
doc.add_paragraph()
doc.add_paragraph()

# Nombre firmante - Centrado, Negrita
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('LIC. YUNERI CALIXTO PÉREZ')
run.bold = True
run.font.size = Pt(10)

# Cargo línea 1 - Centrado
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ENCARGADA DEL DESPACHO DE LA DIRECCIÓN')
run.font.size = Pt(9)

# Cargo línea 2 - Centrado
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DE LA UNIVERSIDAD PEDAGÓGICA NACIONAL')
run.font.size = Pt(9)

# Cargo línea 3 - Centrado
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('UNIDAD 212 TEZIUTLÁN')
run.font.size = Pt(9)

# Guardar documento
doc.save('oficio_template.docx')
print('Plantilla creada exitosamente')
