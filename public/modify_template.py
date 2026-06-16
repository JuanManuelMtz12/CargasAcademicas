from docx import Document

# Abrir el documento existente
doc = Document('oficio_template.docx')

# Función para reemplazar texto en párrafos
def replace_text_in_paragraph(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        # Guardar el formato del primer run
        inline = paragraph.runs
        for run in inline:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)

# Reemplazos en todo el documento
replacements = {
    'LIC. MARLEN ANGEL INOCENCIO': '{teacherName}',
    '10 de septiembre de 2025': '{fecha}',
    'PERSONAL DOCENTE INVITADO DE LA UPN-212 TEZIUTLÁN': '{personalType} DE LA UPN-212 TEZIUTLÁN',
    'LICENCIATURA EN INTERVENCIÓN EDUCATIVA': '{licenciatura}',
    'del 11 de agosto al 5 de diciembre del 2025': '{periodo}'
}

# Reemplazar en párrafos
for paragraph in doc.paragraphs:
    for old, new in replacements.items():
        replace_text_in_paragraph(paragraph, old, new)

# Reemplazar en tablas
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for old, new in replacements.items():
                    replace_text_in_paragraph(paragraph, old, new)

# Modificar la tabla para usar bucle de docxtemplater
# La tabla debería tener el formato correcto para docxtemplater
if len(doc.tables) > 0:
    table = doc.tables[0]
    # Eliminar filas de datos existentes (mantener solo encabezado)
    while len(table.rows) > 1:
        table._element.remove(table.rows[-1]._element)
    
    # Agregar una fila de plantilla
    row = table.add_row()
    row.cells[0].text = '{#horarios}{clave}{/horarios}'
    row.cells[1].text = '{#horarios}{asignatura}{/horarios}'
    row.cells[2].text = '{#horarios}{horario}{/horarios}'
    row.cells[3].text = '{#horarios}{semestre}{/horarios}'

# Guardar el documento modificado
doc.save('oficio_template_modified.docx')
print('Plantilla modificada exitosamente')
